import base64
import logging
import inspect
import random
import pytest

import os
import shutil

from faker import Faker
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook, load_workbook

from Config.config import *


class RandomNumberGenerator:
    def generate_random_number(self):
        random_number = "".join([str(random.randint(0, 9)) for _ in range(10)])
        return random_number


class NameGenerator:
    names = []

    def __init__(self):
        self.fake = Faker()
        # self.names = []

    def generate_name(self):
        name = self.fake.name()
        self.names.append(name)
        return name


class FolderCreator:
    def create_folder(self):
        for i in range(1, 23):
            directory = f"C:\\python-appium\\simple-demo\\Screenshot\\TC_{str(i)}"
            shutil.rmtree(directory)
            os.makedirs(directory)


class DocumentCreator:
    def gen_documentation(self, start, end, folder_num):
        workbook = load_workbook(filename=test_design_dir)
        sheet = workbook.active

        document = Document()
        a = document.add_heading('Simple', 0)
        a.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(start, end):
            document.add_paragraph(sheet[f'A{i}'].value)
            document.add_picture(f'C:\\python-appium\\simple-demo\\Screenshot\\TC_{folder_num}\\'
                                 f'{(str(i - (start - 1)))}.png', width=Inches(1.5))

        document.save(f'C:\\python-appium\\simple-demo\\Screenshot\\TC_{folder_num}\\test.docx')

class LogFunc:
    def get_log(logLevel=logging.DEBUG):

        logger_name = inspect.stack()[1][3]
        logger = logging.getLogger(logger_name)
        logger.setLevel(logging.DEBUG)
        fh = logging.FileHandler(log_dir)
        formatter = logging.Formatter('%(asctime)s - %(levelname)s : %(message)s' , datefmt='%m/%d/%Y %I:%M:%S %p')
        fh.setFormatter(formatter)
        logger.addHandler(fh)
        return logger
